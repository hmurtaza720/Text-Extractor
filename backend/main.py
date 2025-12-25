from fastapi import FastAPI, Depends, HTTPException, status, File, UploadFile, BackgroundTasks
from pydantic import BaseModel
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from . import models, schemas, database, auth
from .tasks import process_document_task
from datetime import datetime
import os
import shutil
import uuid
import io
from xhtml2pdf import pisa
from bs4 import BeautifulSoup
from docx import Document as DocxDocument

# Create DB tables (simplest migration strategy for now)
models.Base.metadata.create_all(bind=database.engine)

# Attempt migration for filename column (simple hack for dev)
from sqlalchemy import text
try:
    with database.engine.connect() as conn:
        conn.execute(text("ALTER TABLE documents ADD COLUMN filename VARCHAR(255)"))
        conn.execute(text("UPDATE documents SET filename = SUBSTRING_INDEX(original_path, '/', -1)")) # Default older
        conn.commit()
except Exception as e:
    pass # Column likely exists

app = FastAPI(title="Secure Document Processor")

# CORS Setup
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], # In production, set to specific frontend origin
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

async def get_current_user(token: str = Depends(oauth2_scheme), db: Session = Depends(database.get_db)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = auth.jwt.decode(token, auth.SECRET_KEY, algorithms=[auth.ALGORITHM])
        user_id: str = payload.get("sub")
        if user_id is None:
            raise credentials_exception
    except auth.JWTError:
        raise credentials_exception
    user = db.query(models.User).filter(models.User.id == int(user_id)).first()
    if user is None:
        raise credentials_exception
    return user

async def get_current_active_user(current_user: models.User = Depends(get_current_user)):
    return current_user

@app.post("/signup", response_model=schemas.UserResponse)
def signup(user: schemas.UserCreate, db: Session = Depends(database.get_db)):
    if user.security_code != auth.INTERNAL_SIGNUP_CODE:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Invalid Security Code. Only internal employees can sign up."
        )

    db_user = db.query(models.User).filter(models.User.email == user.email).first()
    if db_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    hashed_password = auth.get_password_hash(user.password)
    new_user = models.User(username=user.username, email=user.email, hashed_password=hashed_password)
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    return new_user

@app.post("/token", response_model=schemas.Token)
def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(database.get_db)):
    # The frontend code sends email in the 'username' field
    user = db.query(models.User).filter(models.User.email == form_data.username).first()
    if not user or not auth.verify_password(form_data.password, user.hashed_password):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    access_token_expires = auth.timedelta(minutes=auth.ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = auth.create_access_token(
        data={"sub": str(user.id)}, expires_delta=access_token_expires 
    )
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/upload_and_convert", status_code=status.HTTP_202_ACCEPTED)
def upload_and_convert(background_tasks: BackgroundTasks, file: UploadFile = File(...), current_user: models.User = Depends(get_current_active_user), db: Session = Depends(database.get_db)):
    file_id = str(uuid.uuid4())
    extension = os.path.splitext(file.filename)[1]
    file_location = f"uploads/{file_id}{extension}"
    
    os.makedirs("uploads", exist_ok=True)
    
    with open(file_location, "wb+") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    new_doc = models.Document(
        user_id=current_user.id,
        upload_date=str(datetime.utcnow()),
        original_path=file_location,
        filename=file.filename, 
        status="Processing"
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)

    background_tasks.add_task(process_document_task, new_doc.id)

    return {"message": "File uploaded and processing started", "document_id": new_doc.id}

@app.get("/documents", response_model=list[schemas.DocumentResponse])
def get_documents(skip: int = 0, limit: int = 100, current_user: models.User = Depends(get_current_active_user), db: Session = Depends(database.get_db)):
    documents = db.query(models.Document).filter(models.Document.user_id == current_user.id).offset(skip).limit(limit).all()
    return documents

@app.get("/documents/{doc_id}", response_model=schemas.DocumentResponse)
def get_document(doc_id: int, current_user: models.User = Depends(get_current_active_user), db: Session = Depends(database.get_db)):
    doc = db.query(models.Document).filter(models.Document.id == doc_id, models.Document.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc

from typing import Optional

class DocumentUpdate(BaseModel):
    corrected_html: Optional[str] = None
    filename: Optional[str] = None

@app.put("/documents/{doc_id}", response_model=schemas.DocumentResponse)
def update_document(doc_id: int, update_data: DocumentUpdate, current_user: models.User = Depends(get_current_active_user), db: Session = Depends(database.get_db)):
    doc = db.query(models.Document).filter(models.Document.id == doc_id, models.Document.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if update_data.corrected_html is not None:
        doc.corrected_html = update_data.corrected_html
    if update_data.filename is not None:
        doc.filename = update_data.filename
        
    db.commit()
    db.refresh(doc)
    return doc

@app.delete("/documents/{doc_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_document(doc_id: int, current_user: models.User = Depends(get_current_active_user), db: Session = Depends(database.get_db)):
    doc = db.query(models.Document).filter(models.Document.id == doc_id, models.Document.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if doc.original_path and os.path.exists(doc.original_path):
        try:
            os.remove(doc.original_path)
        except OSError as e:
            print(f"Error checking file: {e}") 
    
    db.delete(doc)
    db.commit()
    return None

# User Profile Routes
@app.get("/users/me", response_model=schemas.UserResponse)
async def read_users_me(current_user: models.User = Depends(get_current_active_user)):
    return current_user

@app.put("/users/me", response_model=schemas.UserResponse)
async def update_user_me(user_update: schemas.UserUpdate, current_user: models.User = Depends(get_current_active_user), db: Session = Depends(database.get_db)):
    if user_update.email and user_update.email != current_user.email:
        db_user = db.query(models.User).filter(models.User.email == user_update.email).first()
        if db_user:
            raise HTTPException(status_code=400, detail="Email already registered")
        current_user.email = user_update.email
    
    if user_update.username:
        current_user.username = user_update.username
        
    if user_update.password:
        current_user.hashed_password = auth.get_password_hash(user_update.password)
        
    db.commit()
    db.refresh(current_user)
    return current_user


# --- Tagging Endpoints ---

@app.post("/tags", response_model=schemas.TagResponse)
def create_tag(tag: schemas.TagCreate, db: Session = Depends(database.get_db)):
    # Check if tag exists
    existing_tag = db.query(models.Tag).filter(models.Tag.name == tag.name).first()
    if existing_tag:
        return existing_tag # Return existing if duplicate to avoid errors
    
    new_tag = models.Tag(name=tag.name, color=tag.color)
    db.add(new_tag)
    db.commit()
    db.refresh(new_tag)
    return new_tag

@app.post("/documents/{doc_id}/tags/{tag_name}", response_model=schemas.DocumentResponse)
def add_tag_to_document(doc_id: int, tag_name: str, current_user: models.User = Depends(get_current_active_user), db: Session = Depends(database.get_db)):
    doc = db.query(models.Document).filter(models.Document.id == doc_id, models.Document.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Find or create tag
    tag = db.query(models.Tag).filter(models.Tag.name == tag_name).first()
    if not tag:
        tag = models.Tag(name=tag_name)
        db.add(tag)
        db.commit()
        db.refresh(tag)
        
    if tag not in doc.tags:
        doc.tags.append(tag)
        db.commit()
        db.refresh(doc)
        
    return doc

@app.delete("/documents/{doc_id}/tags/{tag_id}", response_model=schemas.DocumentResponse)
def remove_tag_from_document(doc_id: int, tag_id: int, current_user: models.User = Depends(get_current_active_user), db: Session = Depends(database.get_db)):
    doc = db.query(models.Document).filter(models.Document.id == doc_id, models.Document.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    tag = db.query(models.Tag).filter(models.Tag.id == tag_id).first()
    if tag and tag in doc.tags:
        doc.tags.remove(tag)
        db.commit()
        db.refresh(doc)
        
    return doc


# Frontend Static Files
frontend_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "frontend")
app.mount("/static", StaticFiles(directory=frontend_path), name="static")

@app.get("/")
async def read_index():
    return FileResponse(os.path.join(frontend_path, "index.html"))

@app.get("/dashboard")
async def read_dashboard():
    return FileResponse(os.path.join(frontend_path, "dashboard.html"))

@app.get("/editor")
async def read_editor():
    return FileResponse(os.path.join(frontend_path, "editor.html"))

@app.get("/settings")
async def read_settings():
    return FileResponse(os.path.join(frontend_path, "settings.html"))

# Mount uploads to serve files
app.mount("/uploads", StaticFiles(directory="uploads"), name="uploads")

@app.get("/export/{doc_id}/pdf")
def export_pdf(doc_id: int, current_user: models.User = Depends(get_current_active_user), db: Session = Depends(database.get_db)):
    doc = db.query(models.Document).filter(models.Document.id == doc_id, models.Document.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    html_content = doc.corrected_html or doc.raw_text or ""
    # Wrap in basic HTML structure if missing
    if "<html>" not in html_content:
        html_content = f"<html><body>{html_content}</body></html>"
        
    pdf_buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(html_content, dest=pdf_buffer)
    
    if pisa_status.err:
        raise HTTPException(status_code=500, detail="Error generating PDF")
        
    pdf_buffer.seek(0)
    filename = (doc.filename or "document").replace(".pdf", "") + ".pdf"
    
    return StreamingResponse(
        pdf_buffer, 
        media_type="application/pdf", 
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@app.get("/export/{doc_id}/docx")
def export_docx(doc_id: int, current_user: models.User = Depends(get_current_active_user), db: Session = Depends(database.get_db)):
    doc = db.query(models.Document).filter(models.Document.id == doc_id, models.Document.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
        
    html_content = doc.corrected_html or doc.raw_text or ""
    soup = BeautifulSoup(html_content, "html.parser")
    
    docx = DocxDocument()
    docx.add_heading(doc.filename or "Document", 0)
    
    # Very basic HTML interpretation
    for element in soup.descendants:
        if element.name in ['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li']:
             text = element.get_text(strip=True)
             if text:
                 if element.name.startswith('h'):
                     try:
                        level = int(element.name[1])
                        docx.add_heading(text, level=level)
                     except:
                        docx.add_paragraph(text)
                 elif element.name == 'li':
                     docx.add_paragraph(text, style='List Bullet')
                 else:
                     docx.add_paragraph(text)
                     
    docx_buffer = io.BytesIO()
    docx.save(docx_buffer)
    docx_buffer.seek(0)
    
    filename = (doc.filename or "document").replace(".docx", "") + ".docx"
    
    return StreamingResponse(
        docx_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

class N8NCallback(BaseModel):
    doc_id: int
    raw_text: str
    corrected_html: Optional[str] = None
    status: str = "Ready"

@app.post("/n8n/callback")
def n8n_callback(data: N8NCallback, db: Session = Depends(database.get_db)):
    doc = db.query(models.Document).filter(models.Document.id == data.doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
        
    doc.raw_text = data.raw_text
    # If html is provided, use it. If not, generate basic wrapper.
    if data.corrected_html:
        doc.corrected_html = data.corrected_html
    else:
        # Simple fallback conversion
        safe_text = data.raw_text.replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br>")
        doc.corrected_html = f"<div>{safe_text}</div>"
        
    doc.status = data.status
    db.commit()
    return {"status": "success"}

